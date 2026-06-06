import docx
import json

def ler_docx(file_path):
    doc = docx.Document(file_path)
    doc_text = "\n".join([para.text for para in doc.paragraphs])
    return doc_text

"""def read_docx_dict(file_path):
    original_doc = docx.Document(file_path)
    doc_dict = []

    for para in original_doc.paragraphs:
        if para.text.strip():  # Only process non-empty paragraphs
            doc_dict.append({
                "alignment": para.alignment,
                "runs": []
            })
            
            for run in para.runs:
                if run.text.strip():  # Only process non-empty runs
                    doc_dict[-1]["runs"].append({
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "font_name": run.font.name,
                        "font_size": run.font.size
                    })

    return doc_dict  

def parse_docx(doc_dict):
    json_doc = json.dumps(doc_dict, indent=4, ensure_ascii=False)
    return json_doc"""